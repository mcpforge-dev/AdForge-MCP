from __future__ import annotations

import base64
import io
from typing import Any


MAX_REPORT_TEXT = 4_000
MAX_REPORT_IMAGE_BYTES = 180_000


def build_site_analysis_docx(analysis: dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:  # pragma: no cover - optional dependency is verified during deploy.
        raise RuntimeError("python-docx is not installed") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for style_name, size, color in (
        ("Title", 24, "111827"),
        ("Heading 1", 17, "111827"),
        ("Heading 2", 13, "24324A"),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = document.add_heading("AI-анализ сайта HolyMedia MCP", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(_text(analysis.get("url")))
    score = document.add_paragraph()
    score_run = score.add_run(f"Оценка сайта: {_text(analysis.get('overall_score'), '—')}/100")
    score_run.bold = True
    score_run.font.size = Pt(16)

    document.add_heading("Краткий вердикт", level=1)
    verdict = analysis.get("verdict", {}) if isinstance(analysis.get("verdict"), dict) else {}
    document.add_paragraph(_text(verdict.get("summary") or analysis.get("summary")))
    _label_paragraph(document, "Главный риск", verdict.get("main_risk"))
    _label_paragraph(document, "Быстрый выигрыш", verdict.get("fastest_win"))

    overview = analysis.get("audit_overview", {}) if isinstance(analysis.get("audit_overview"), dict) else {}
    confidence = overview.get("confidence", {}) if isinstance(overview.get("confidence"), dict) else {}
    document.add_heading("Диагностика и доказательства", level=1)
    _label_paragraph(document, "Достоверность", f"{_text(confidence.get('score'), '0')}% · {_text(confidence.get('label'), 'ограниченная')}")
    _label_paragraph(document, "Источники", ", ".join(_string_list(confidence.get("sources"))))
    _add_screenshots(document, overview.get("screenshots", {}))

    diagnostic_rows: list[list[str]] = []
    for pillar in _dict_list(overview.get("pillars")):
        for check in _dict_list(pillar.get("checks")):
            diagnostic_rows.append(
                [
                    _text(pillar.get("title")),
                    f"{_text(pillar.get('score'), '—')}/100",
                    _text(check.get("title")),
                    _status_label(check.get("status")),
                    _text(check.get("evidence")),
                    "—" if check.get("status") == "pass" else _text(check.get("action")),
                ]
            )
    _add_table(document, ["Направление", "Оценка", "Проверка", "Статус", "Evidence", "Действие"], diagnostic_rows)

    document.add_heading("Топ улучшений", level=1)
    issue_rows = [
        [
            _text(item.get("priority"), "P2"),
            _text(item.get("title")),
            _text(item.get("problem")),
            _text(item.get("what_to_do")),
            _text(item.get("evidence")),
        ]
        for item in _dict_list(analysis.get("top_issues"))
    ]
    _add_table(document, ["Приоритет", "Проблема", "Диагноз", "Что сделать", "Evidence"], issue_rows)

    first_screen = analysis.get("first_screen_review", {}) if isinstance(analysis.get("first_screen_review"), dict) else {}
    found = first_screen.get("found", {}) if isinstance(first_screen.get("found"), dict) else {}
    example = first_screen.get("example_hero", {}) if isinstance(first_screen.get("example_hero"), dict) else {}
    document.add_heading("Разбор первого экрана", level=1)
    document.add_paragraph(_text(first_screen.get("five_second_takeaway")))
    _label_paragraph(document, "Найденный H1", found.get("h1") or "не найден")
    _label_paragraph(document, "Найденные CTA", ", ".join(_string_list(found.get("ctas"))) or "не найдены")
    for item in _string_list(first_screen.get("friction")):
        document.add_paragraph(item, style="List Bullet")
    document.add_heading(_text(example.get("label"), "Пример первого экрана"), level=2)
    _label_paragraph(document, "H1", example.get("h1"))
    _label_paragraph(document, "Подзаголовок", example.get("subtitle"))
    _label_paragraph(document, "Основной CTA", example.get("primary_cta"))
    _label_paragraph(document, "Визуальное направление", example.get("visual_direction"))

    document.add_heading("Что сделать за 1 день", level=1)
    day_rows = [
        [
            _text(item.get("task")),
            _text(item.get("owner")),
            _text(item.get("time")),
            _text(item.get("expected_effect")),
            _text(item.get("placement")),
        ]
        for item in _dict_list(analysis.get("one_day_plan"))
    ]
    _add_table(document, ["Задача", "Ответственный", "Время", "Эффект", "Где внедрить"], day_rows)

    document.add_heading("План внедрения", level=1)
    plan_rows = [
        [
            _text(item.get("task")),
            _text(item.get("impact")),
            _text(item.get("difficulty")),
            _text(item.get("priority")),
            _text(item.get("owner")),
        ]
        for item in _dict_list(analysis.get("implementation_plan"))
    ]
    _add_table(document, ["Задача", "Влияние", "Сложность", "Приоритет", "Ответственный"], plan_rows)

    document.add_heading("Границы анализа", level=1)
    for item in _string_list(confidence.get("limitations")):
        document.add_paragraph(item, style="List Bullet")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _add_screenshots(document: Any, screenshots: Any) -> None:
    from docx.shared import Inches

    if not isinstance(screenshots, dict):
        return
    for label, key, width in (("Desktop screenshot", "desktop", 6.4), ("Mobile screenshot", "mobile", 3.0)):
        item = screenshots.get(key, {})
        if not isinstance(item, dict):
            continue
        content = _decode_image(str(item.get("preview_data_url", "")))
        if not content:
            continue
        document.add_heading(label, level=2)
        document.add_picture(io.BytesIO(content), width=Inches(width))


def _decode_image(data_url: str) -> bytes:
    if not data_url.startswith("data:image/jpeg;base64,"):
        return b""
    try:
        content = base64.b64decode(data_url.partition(",")[2], validate=True)
    except (ValueError, TypeError):
        return b""
    return content if len(content) <= MAX_REPORT_IMAGE_BYTES else b""


def _add_table(document: Any, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        document.add_paragraph("Нет данных для этого раздела.")
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row[: len(headers)]):
            cells[index].text = _text(value)


def _label_paragraph(document: Any, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(_text(value))


def _dict_list(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def _string_list(value: Any) -> list[str]:
    return [_text(item) for item in value if _text(item)] if isinstance(value, list) else []


def _status_label(value: Any) -> str:
    labels = {"pass": "пройдено", "warn": "нужно проверить", "fail": "высокий риск", "unknown": "нет данных"}
    return labels.get(str(value), "нет данных")


def _text(value: Any, fallback: str = "") -> str:
    text = " ".join(str(value if value is not None else fallback).split())
    return text[:MAX_REPORT_TEXT]
