from __future__ import annotations

import io
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    Document = None


NAVY = "14213D"
BLUE = "2563EB"
PALE_BLUE = "EAF1FF"
PALE_GRAY = "F4F6F8"
MID_GRAY = "667085"
TEXT = "1F2937"
GREEN = "117A65"
AMBER = "9A6700"


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, color: str = "D0D5DD", size: str = "6") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _set_cell_text(cell: Any, text: str, *, color: str = TEXT, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table(document: Any, headers: list[str], rows: list[list[str]], *, widths: list[float] | None = None) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, color="FFFFFF", bold=True, size=8)
        _set_cell_shading(table.rows[0].cells[index], NAVY)
        _set_cell_border(table.rows[0].cells[index], NAVY)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value, size=8)
            _set_cell_shading(cells[index], "FFFFFF" if len(table.rows) % 2 else PALE_GRAY)
            _set_cell_border(cells[index])
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _add_heading(document: Any, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Aptos Display"
    run.font.bold = True
    run.font.size = Pt(17 if level == 1 else 12)
    run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)


def _add_body(document: Any, text: str, *, color: str = TEXT, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _fmt(value: Any, *, currency: str = "", percent: bool = False) -> str:
    if value is None:
        return "Данные не предоставлены"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if percent:
        return f"{number:.2f}%"
    if abs(number - round(number)) < 0.000001:
        result = f"{int(round(number)):,}".replace(",", " ")
    else:
        result = f"{number:,.2f}".replace(",", " ")
    return f"{result} {currency}".strip()


def _metric(dataset: dict[str, Any], name: str) -> Any:
    item = dataset.get("metrics", {}).get(name, {})
    return item.get("value") if isinstance(item, dict) else None


def _add_callout(document: Any, title: str, text: str, *, fill: str = PALE_BLUE, color: str = BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_border(cell, fill, "0")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    title_run = paragraph.add_run(title + "\n")
    title_run.font.name = "Aptos"
    title_run.font.size = Pt(9)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(color)
    text_run = paragraph.add_run(text)
    text_run.font.name = "Aptos"
    text_run.font.size = Pt(9)
    text_run.font.color.rgb = RGBColor.from_string(TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def build_monthly_ads_report_docx(dataset: dict[str, Any]) -> bytes:
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx is not installed")

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXT)

    account = dataset.get("account", {})
    provider = str(dataset.get("provider", "ads")).replace("_", " ").title()
    period = dataset.get("period", {})
    currency = str(account.get("currency") or "USD")
    source = dataset.get("source", {})

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover.paragraph_format.space_before = Pt(40)
    run = cover.add_run("HOLYMEDIA MCP")
    run.font.name = "Aptos Display"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Месячный отчёт\nпо рекламе")
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    _add_body(document, f"{account.get('name', account.get('account_id', 'Рекламный кабинет'))} · {provider}", color=MID_GRAY, bold=True)
    _add_body(document, f"Период: {period.get('start', '—')} — {period.get('end', '—')} · Часовой пояс: {period.get('timezone', '—')}", color=MID_GRAY)
    _add_callout(
        document,
        "Статус данных",
        "Живые данные рекламного кабинета" if source.get("real_data") else "Источник не вернул полный набор живых строк; пропуски отмечены в отчёте.",
        fill="EAF7F2" if source.get("real_data") else "FFF4D6",
        color=GREEN if source.get("real_data") else AMBER,
    )
    document.add_page_break()

    _add_heading(document, "1. Общие итоги")
    _add_body(document, "Ниже приведены только показатели, полученные из подключённого рекламного аккаунта или рассчитанные из них. Отсутствующие значения не заменяются нулём.")
    kpi_rows = [
        ["Расход", _fmt(_metric(dataset, "spend"), currency=currency)],
        ["Показы", _fmt(_metric(dataset, "impressions"))],
        ["Клики", _fmt(_metric(dataset, "clicks"))],
        ["Конверсии платформы", _fmt(_metric(dataset, "conversions"))],
        ["CTR", _fmt(_metric(dataset, "ctr"), percent=True)],
        ["Средняя стоимость клика", _fmt(_metric(dataset, "cpc"), currency=currency)],
        ["Стоимость конверсии", _fmt(_metric(dataset, "cpa"), currency=currency)],
    ]
    _add_table(document, ["Показатель", "Значение"], kpi_rows, widths=[3.7, 2.7])

    _add_heading(document, "2. Сравнение с предыдущим периодом")
    comparison = dataset.get("comparison", {})
    comparison_rows: list[list[str]] = []
    labels = {"spend": "Расход", "impressions": "Показы", "clicks": "Клики", "conversions": "Конверсии", "ctr": "CTR", "cpc": "CPC", "cpa": "CPA"}
    for metric, label in labels.items():
        item = comparison.get(metric, {})
        comparison_rows.append([
            label,
            _fmt(item.get("current"), currency=currency, percent=metric == "ctr"),
            _fmt(item.get("previous"), currency=currency, percent=metric == "ctr"),
            _fmt(item.get("delta_percent"), percent=True),
        ])
    _add_table(document, ["Показатель", "Текущий", "Предыдущий", "Изменение"], comparison_rows, widths=[2.3, 1.55, 1.55, 1.3])

    _add_heading(document, "3. Подтверждённые выводы")
    assertions = dataset.get("assertions", [])
    if assertions:
        for assertion in assertions:
            _add_body(document, f"• {assertion.get('text', '')}")
    else:
        _add_callout(document, "Недостаточно данных", "Источник не вернул фактов, на которых можно построить подтверждённый вывод.", fill="FFF4D6", color=AMBER)

    _add_heading(document, "4. Структура кампаний")
    campaigns = dataset.get("campaigns", [])
    campaign_rows = []
    for campaign in campaigns[:20]:
        campaign_rows.append([
            str(campaign.get("entity_name") or "Без названия"),
            _fmt(campaign.get("spend"), currency=currency),
            _fmt(campaign.get("impressions")),
            _fmt(campaign.get("clicks")),
            _fmt(campaign.get("conversions")),
            _fmt(campaign.get("spend_share_percent"), percent=True),
        ])
    if campaign_rows:
        _add_table(document, ["Кампания", "Расход", "Показы", "Клики", "Конверсии", "Доля"], campaign_rows, widths=[2.55, 0.9, 0.85, 0.75, 0.9, 0.75])
    else:
        _add_callout(document, "Кампании не показаны", "За период нет строк уровня campaign или источник вернул пустой ответ.", fill="FFF4D6", color=AMBER)

    _add_heading(document, "5. Ключевые события и журнал изменений")
    changes = dataset.get("changes", [])
    if changes:
        change_rows = [[str(item.get("changed_at", "—")), str(item.get("entity_name") or item.get("entity_id") or "—"), str(item.get("field", "—")), str(item.get("reason_status", "—"))] for item in changes]
        _add_table(document, ["Дата", "Объект", "Поле", "Причина"], change_rows, widths=[1.35, 2.2, 1.35, 1.5])
    else:
        _add_callout(document, "Изменения не подтверждены", "Журнал изменений не был передан в этот запуск. Причины динамики в отчёте не утверждаются.", fill="FFF4D6", color=AMBER)

    _add_heading(document, "6. Креативы, вовлечение и дополнительные данные")
    pending = dataset.get("pending_sections", [])
    pending_text = ", ".join(pending) if pending else "Нет"
    _add_callout(document, "Данные не предоставлены", f"В текущий запуск не входят: {pending_text}.", fill="FFF4D6", color=AMBER)
    _add_body(document, "Эти разделы не заполняются догадками и не заменяются агрегированными значениями из другого источника.")

    _add_heading(document, "7. Рекомендации")
    for recommendation in dataset.get("recommendations", []):
        _add_body(document, f"• {recommendation.get('text', '')}")
        condition = recommendation.get("condition")
        if condition:
            _add_body(document, f"Условие: {condition}", color=MID_GRAY)

    _add_heading(document, "8. Аналитика и ограничения")
    for limitation in dataset.get("limitations", []):
        _add_body(document, f"• {limitation}")
    unknowns = dataset.get("unknowns", [])
    if unknowns:
        _add_callout(document, "Показатели без достоверного значения", ", ".join(unknowns), fill="FFF4D6", color=AMBER)

    _add_heading(document, "9. Вопросы клиенту")
    questions = dataset.get("questions", [])
    question_rows = [[str(item.get("question_id", "")), str(item.get("question", ""))] for item in questions]
    if question_rows:
        _add_table(document, ["ID", "Что нужно уточнить"], question_rows, widths=[1.35, 5.05])
    else:
        _add_body(document, "Открытых вопросов нет.")

    _add_heading(document, "10. Реестр источников")
    source_rows = [[str(source.get("source_api", "—")), str(source.get("data_status", "—")), str(source.get("fetched_at", "—"))]]
    _add_table(document, ["Источник", "Статус", "Получено"], source_rows, widths=[2.7, 1.4, 2.3])
    _add_body(document, f"Report run ID: {dataset.get('report_run_id', '—')}", color=MID_GRAY)
    _add_body(document, "Отчёт сформирован HolyMedia MCP. Числа относятся только к выбранному рекламному кабинету и указанному периоду.", color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("HolyMedia MCP · проверенный отчёт")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(MID_GRAY)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
